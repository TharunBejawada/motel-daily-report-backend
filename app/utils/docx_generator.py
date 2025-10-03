# app/utils/docx_generator.py
from docx import Document
from docx.shared import Pt, Inches


def _h1(doc: Document, text: str):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.size = Pt(20)
    return p


def _kv_table(doc: Document, title: str, rows: list[list[str]]):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(rows[0]))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(rows[0]):
        hdr_cells[i].text = str(h)
    for r in rows[1:]:
        row_cells = table.add_row().cells
        for i, c in enumerate(r):
            row_cells[i].text = "" if c is None else str(c)
    doc.add_paragraph("")


def build_report_docx(buf, data: dict):
    doc = Document()
    _h1(doc, "Daily Report")

    _kv_table(doc, "Overview", [
        ["Field", "Value"],
        ["Motel", data.get("motel_name") or data.get("property_name", "")],
        ["Location", data.get("location") or ""],
        ["Report Date", data.get("report_date") or ""],
        ["Department", data.get("department") or ""],
        ["Auditor", data.get("auditor") or ""],
    ])

    _kv_table(doc, "KPIs", [
        ["Metric", "Value"],
        ["Revenue", data.get("revenue")],
        ["ADR", data.get("adr")],
        ["Occupancy", data.get("occupancy")],
        ["Vacant Clean", data.get("vacant_clean")],
        ["Vacant Dirty", data.get("vacant_dirty")],
        ["Out of Order / Storage Rooms", data.get("out_of_order_storage_rooms")],
    ])

    if data.get("comp_rooms"):
        _kv_table(doc, "Complementary Rooms", [
            ["Room #", "Notes"],
            *[[r.get("room_number", ""), r.get("notes", "")] for r in data["comp_rooms"]],
        ])

    if data.get("vacant_dirty_rooms"):
        _kv_table(doc, "Vacant / Dirty Rooms", [
            ["Room #", "Reason", "Days", "Action"],
            *[[r.get("room_number", ""), r.get("reason", ""), r.get("days", ""), r.get("action", "")]
              for r in data["vacant_dirty_rooms"]],
        ])

    if data.get("out_of_order_rooms"):
        _kv_table(doc, "Out of Order / Storage", [
            ["Room #", "Reason", "Days", "Action"],
            *[[r.get("room_number", ""), r.get("reason", ""), r.get("days", ""), r.get("action", "")]
              for r in data["out_of_order_rooms"]],
        ])

    if data.get("incidents"):
        _kv_table(doc, "Incidents", [
            ["Description"],
            *[[r.get("description", "")] for r in data["incidents"]],
        ])

    doc.save(buf)
